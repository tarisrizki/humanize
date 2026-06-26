"""HumanWrite AI — Enhanced Streamlit GUI (Global Pre-trained Mode).

Provides a premium interface for humanizing LLM-generated drafts using a
pre-trained global style profile.
"""

import difflib
import os
import io
import time
import json

import requests
import streamlit as st
import pandas as pd

# ── Page Configuration ────────────────────────────────────────────────────────

st.set_page_config(
    page_title="HumanWrite AI",
    page_icon="✍️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Custom CSS for Premium Look ───────────────────────────────────────────────

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');

    /* Global */
    .stApp {
        font-family: 'Inter', sans-serif;
    }

    /* Header */
    .main-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 2rem 2.5rem;
        border-radius: 16px;
        margin-bottom: 2rem;
        color: white;
        box-shadow: 0 10px 40px rgba(102, 126, 234, 0.3);
    }
    .main-header h1 {
        margin: 0;
        font-size: 2rem;
        font-weight: 700;
        letter-spacing: -0.5px;
    }
    .main-header p {
        margin: 0.5rem 0 0 0;
        opacity: 0.9;
        font-size: 1rem;
        font-weight: 300;
    }

    /* Metric cards */
    .metric-card {
        background: linear-gradient(145deg, #1e1e2e, #2a2a3e);
        border: 1px solid rgba(255,255,255,0.08);
        border-radius: 12px;
        padding: 1.2rem;
        text-align: center;
        transition: transform 0.2s ease, box-shadow 0.2s ease;
    }
    .metric-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 25px rgba(0,0,0,0.3);
    }
    .metric-label {
        font-size: 0.75rem;
        text-transform: uppercase;
        letter-spacing: 1px;
        color: #a0a0b0;
        margin-bottom: 0.4rem;
    }
    .metric-value {
        font-size: 1.6rem;
        font-weight: 700;
        color: #667eea;
    }

    /* Diff view */
    .diff-add {
        background-color: rgba(46, 160, 67, 0.15);
        color: #3fb950;
        padding: 2px 4px;
        border-radius: 3px;
    }
    .diff-remove {
        background-color: rgba(248, 81, 73, 0.15);
        color: #f85149;
        text-decoration: line-through;
        padding: 2px 4px;
        border-radius: 3px;
    }
    .diff-container {
        background: #0d1117;
        border: 1px solid #30363d;
        border-radius: 8px;
        padding: 1rem;
        font-family: 'Consolas', 'Monaco', monospace;
        font-size: 0.85rem;
        line-height: 1.6;
        overflow-x: auto;
    }

    /* Score gauge */
    .gauge-container {
        text-align: center;
        padding: 1rem;
    }
    .gauge-score {
        font-size: 3rem;
        font-weight: 700;
        background: linear-gradient(135deg, #667eea, #764ba2);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    .gauge-label {
        font-size: 0.85rem;
        color: #a0a0b0;
        text-transform: uppercase;
        letter-spacing: 1px;
    }

    /* Status indicator */
    .status-dot {
        display: inline-block;
        width: 8px;
        height: 8px;
        border-radius: 50%;
        margin-right: 6px;
    }
    .status-online { background-color: #3fb950; }
    .status-offline { background-color: #f85149; }

    /* Sidebar styling */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #0d1117 0%, #161b22 100%);
    }

    /* Changes list */
    .change-item {
        background: rgba(102, 126, 234, 0.08);
        border-left: 3px solid #667eea;
        padding: 0.6rem 1rem;
        margin: 0.4rem 0;
        border-radius: 0 6px 6px 0;
        font-size: 0.9rem;
    }
</style>
""", unsafe_allow_html=True)


# ── Configuration ─────────────────────────────────────────────────────────────

BACKEND_URL = os.environ.get("BACKEND_URL", "http://127.0.0.1:8000")


def check_backend() -> bool:
    """Check if the backend is running."""
    try:
        r = requests.get(f"{BACKEND_URL}/health", timeout=3)
        return r.status_code == 200
    except Exception:
        return False

def get_global_style() -> dict:
    """Fetch the pre-trained global style profile."""
    try:
        r = requests.get(f"{BACKEND_URL}/api/v1/style", timeout=5)
        if r.status_code == 200:
            return r.json()
    except Exception:
        pass
    return {}

def process_stream(response):
    """Generator to process the SSE streaming response from the backend."""
    current_event = None
    for line in response.iter_lines(decode_unicode=True):
        if not line:
            continue
            
        if line.startswith("event: "):
            current_event = line[len("event: "):].strip()
        elif line.startswith("data: "):
            data_str = line[len("data: "):]
            if current_event == "text":
                try:
                    chunk = json.loads(data_str)
                    yield chunk
                except:
                    pass
            elif current_event == "metrics":
                try:
                    st.session_state["last_metrics"] = json.loads(data_str)
                except json.JSONDecodeError:
                    pass

# ── Header ────────────────────────────────────────────────────────────────────

st.markdown("""
<div class="main-header">
    <h1>✍️ HumanWrite AI</h1>
    <p>Voice Preservation Engine — Pre-trained Global Human Style</p>
</div>
""", unsafe_allow_html=True)


# ── Sidebar ───────────────────────────────────────────────────────────────────

with st.sidebar:
    st.markdown("### ⚙️ System Status")

    # Backend status
    backend_online = check_backend()
    status_class = "status-online" if backend_online else "status-offline"
    status_text = "Backend Online" if backend_online else "Backend Offline"
    st.markdown(
        f'<span class="status-dot {status_class}"></span> {status_text}',
        unsafe_allow_html=True,
    )

    if not backend_online:
        st.warning("⚠️ Backend is not running.")
        
    st.markdown("---")
    st.markdown(
        "<small style='color: #666;'>HumanWrite AI v0.2.0<br>"
        "Powered by Groq LLaMA 3.3-70b</small>",
        unsafe_allow_html=True,
    )


# ── Main Content ──────────────────────────────────────────────────────────────

profile = get_global_style() if backend_online else {}

if profile:
    with st.expander("📊 Lihat Metrik Profil Global", expanded=False):
        st.markdown(f"**Bahasa Utama:** `{profile.get('language', 'id').upper()}`")
        
        # Metric cards
        cols = st.columns(4)
        with cols[0]:
            st.markdown(f"""<div class="metric-card"><div class="metric-label">Rata-rata Kata/Kalimat</div><div class="metric-value">{profile.get('avg_sentence_length', 0):.1f}</div></div>""", unsafe_allow_html=True)
        with cols[1]:
            st.markdown(f"""<div class="metric-card"><div class="metric-label">Readability (Flesch)</div><div class="metric-value">{profile.get('flesch_reading_ease', 0):.1f}</div></div>""", unsafe_allow_html=True)
        with cols[2]:
            st.markdown(f"""<div class="metric-card"><div class="metric-label">Kalimat Aktif</div><div class="metric-value">{profile.get('active_voice_ratio', 0):.0%}</div></div>""", unsafe_allow_html=True)
        with cols[3]:
            st.markdown(f"""<div class="metric-card"><div class="metric-label">Register</div><div class="metric-value" style="font-size:1.2rem;">{profile.get('emotion_register', 'neutral').title()}</div></div>""", unsafe_allow_html=True)

if not profile and backend_online:
    st.error("Global Style Profile tidak ditemukan. Harap jalankan script training di backend.")

tab1, tab2, tab3 = st.tabs([
    "✍️ Humanize", 
    "📊 Evaluasi", 
    "📜 History"
])

# ── Tab 1: Humanize ───────────────────────────────────────────────────────────
with tab1:
    st.markdown("### Masukkan Teks AI")
    st.caption("Tempel draf yang dihasilkan AI di sini untuk ditulis ulang menjadi lebih natural.")

    style_mode_options = ["populer", "akademik", "profesional", "kreatif"]
    style_mode_val = st.selectbox(
        "Pilih Mode Gaya Penulisan",
        options=style_mode_options,
        index=0,
    )

    draft_text = st.text_area(
        "Teks Draf",
        height=250,
        placeholder="Tempel draf AI Anda di sini...",
        label_visibility="collapsed",
    )

    humanize_btn = st.button(
        "🪄 Humanize Teks",
        use_container_width=True,
        disabled=not draft_text or not backend_online or not profile,
    )

    if humanize_btn and draft_text:
        with st.spinner("✨ Sedang menulis ulang dengan gaya manusia..."):
            try:
                process_resp = requests.post(
                    f"{BACKEND_URL}/api/v1/process",
                    json={"draft": draft_text, "style_mode": style_mode_val},
                    stream=True,
                    timeout=120,
                )
                process_resp.raise_for_status()
                
                st.markdown("### ✅ Hasil Humanize")
                st.session_state["last_metrics"] = None
                
                # Create a placeholder for the streamed text
                text_placeholder = st.empty()
                
                # Consume the stream
                final_text = text_placeholder.write_stream(process_stream(process_resp))
                
                st.session_state["last_result"] = final_text
                st.session_state["last_draft"] = draft_text
                st.session_state["style_mode_val"] = style_mode_val
                
                # Log the evaluation to get a record_id
                eval_resp = requests.post(
                    f"{BACKEND_URL}/api/v1/evaluate/run",
                    json={
                        "style_mode": style_mode_val,
                        "language": profile.get("language", "id"),
                        "original_text": draft_text,
                        "output_text": final_text,
                    },
                    timeout=10,
                )
                if eval_resp.status_code == 200:
                    st.session_state["last_record_id"] = eval_resp.json().get("record_id")
                
            except Exception as e:
                st.error(f"Proses gagal: {e}")

    # Display results
    if st.session_state.get("last_result"):
        final_text = st.session_state["last_result"]
        original_draft = st.session_state.get("last_draft", "")
        metrics = st.session_state.get("last_metrics") or {}
        changes = metrics.get("changes_made", [])

        st.markdown("---")
        
        # Word count
        word_count_orig = len(original_draft.split())
        word_count_final = len(final_text.split())
        st.caption(f"**Jumlah Kata:** Draf Asli ({word_count_orig}) ➡️ Hasil Humanize ({word_count_final})")

        # ── Diff View ─────────────────────────────────────────────────
        with st.expander("🔀 Tampilan Diff (perubahan baris per baris)", expanded=True):
            diff = difflib.unified_diff(
                original_draft.splitlines(keepends=True),
                final_text.splitlines(keepends=True),
                fromfile="Draf Asli",
                tofile="Teks Hasil Humanize",
                lineterm="",
            )
            diff_lines = list(diff)

            if diff_lines:
                html_parts = []
                for line in diff_lines:
                    line_escaped = (
                        line.replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;")
                    )
                    if line.startswith("+") and not line.startswith("+++"):
                        html_parts.append(f'<span class="diff-add">{line_escaped}</span>')
                    elif line.startswith("-") and not line.startswith("---"):
                        html_parts.append(f'<span class="diff-remove">{line_escaped}</span>')
                    elif line.startswith("@@"):
                        html_parts.append(f'<span style="color:#8b949e;">{line_escaped}</span>')
                    else:
                        html_parts.append(line_escaped)

                diff_html = "<br>".join(html_parts)
                st.markdown(
                    f'<div class="diff-container">{diff_html}</div>',
                    unsafe_allow_html=True,
                )
            else:
                st.info("Tidak ada perbedaan signifikan yang terdeteksi.")

        # ── Changes List ──────────────────────────────────────────────
        if changes:
            with st.expander("📋 Perubahan yang Dilakukan", expanded=True):
                for i, change in enumerate(changes, 1):
                    st.markdown(
                        f'<div class="change-item">💡 {change}</div>',
                        unsafe_allow_html=True,
                    )

        # ── Export Buttons ────────────────────────────────────────────
        st.markdown("---")
        st.markdown("#### 📥 Ekspor & Salin")
        export_col1, export_col2, export_col3 = st.columns(3)

        with export_col1:
            st.download_button(
                label="📄 Download .txt",
                data=final_text,
                file_name=f"humanized_global_{int(time.time())}.txt",
                mime="text/plain",
                use_container_width=True,
            )

        with export_col2:
            try:
                from docx import Document as DocxDocument

                doc = DocxDocument()
                doc.add_heading("HumanWrite AI — Hasil Humanize", level=1)
                doc.add_paragraph("Style: Global Pre-trained Model")
                doc.add_heading("Teks Hasil Humanize", level=2)
                for para in final_text.split("\n\n"):
                    doc.add_paragraph(para)

                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.download_button(
                    label="📝 Download .docx",
                    data=buffer,
                    file_name=f"humanized_global_{int(time.time())}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except ImportError:
                st.button("📝 .docx (install python-docx dahulu)", disabled=True, use_container_width=True)

        with export_col3:
            st.code(final_text, language='text')

# ── Tab 2: Evaluasi ───────────────────────────────────────────────────────────
with tab2:
    st.markdown("### 🤖 Evaluasi LLM Judge & Anti-Deteksi")
    if "last_record_id" in st.session_state and st.session_state.get("last_result"):
        original_draft = st.session_state.get("last_draft", "")
        final_text = st.session_state.get("last_result", "")
        
        # LLM Judge
        st.markdown("#### ⚖️ LLM as a Judge")
        st.caption("Evaluasi kualitas hasil humanize menggunakan penalaran LLM Llama-3.3.")
        if st.button("🤖 Jalankan Evaluasi Otomatis", use_container_width=True):
            with st.spinner("Llama-3.3-70b sedang mengevaluasi kualitas tulisan..."):
                try:
                    resp = requests.post(
                        f"{BACKEND_URL}/api/v1/evaluate/judge",
                        json={
                            "record_id":      st.session_state["last_record_id"],
                            "original_text":  original_draft,
                            "humanized_text": final_text,
                            "style_mode":     st.session_state.get("style_mode_val", "populer"),
                            "language":       profile.get("language", "id"),
                        },
                        timeout=60
                    )
                    resp.raise_for_status()
                    data = resp.json()
                    
                    score = data["overall_score"]
                    color = "green" if score >= 70 else "orange" if score >= 50 else "red"
                    
                    st.markdown(f"### Skor Keseluruhan: :{color}[{score}/100]")
                    
                    col_j1, col_j2 = st.columns(2)
                    breakdown = data["breakdown"]
                    
                    with col_j1:
                        for dim in ["naturalness", "register_compliance", "content_fidelity"]:
                            d = breakdown[dim]
                            st.metric(dim.replace("_", " ").title(), f"{d['score']}/10")
                            st.caption(d["reason"])
                            
                    with col_j2:
                        for dim in ["eyd_grammar", "anti_detection"]:
                            d = breakdown[dim]
                            st.metric(dim.replace("_", " ").title(), f"{d['score']}/10")
                            st.caption(d["reason"])
                            
                    if breakdown["critical_issues"]:
                        st.error("⚠️ Isu Kritis:\n" + "\n".join(f"• {i}" for i in breakdown["critical_issues"]))
                        
                    st.success(f"✅ Sorotan Terbaik: *{breakdown['highlight']}*")
                    st.warning(f"⚠️ Paling AI: *{breakdown['worst_sentence']}*")
                except Exception as e:
                    st.error(f"Gagal menjalankan LLM Judge: {e}")
        # Turnitin Safety
        st.markdown("---")
        st.markdown("#### 🛡️ Turnitin Safety")

        if st.session_state.get("last_result") and st.session_state.get("last_draft"):
            # Cari info overlap dari changes_made
            changes = st.session_state.get("last_metrics", {}).get("changes_made", [])
            overlap_info = next(
                (c for c in changes if "Kesamaan struktural" in c), None
            )

            if overlap_info:
                # Parse persentase dari string
                import re
                pct_match = re.search(r'(\d+)%', overlap_info)
                if pct_match:
                    overlap_pct = int(pct_match.group(1))

                    if overlap_pct < 15:
                        st.success(f"✅ **{overlap_pct}% kesamaan struktural** — Aman untuk Turnitin")
                    elif overlap_pct < 30:
                        st.warning(f"⚠️ **{overlap_pct}% kesamaan struktural** — Perlu perhatian")
                    else:
                        st.error(f"🔴 **{overlap_pct}% kesamaan struktural** — Risiko tinggi plagiarisme")

                    # Progress bar visual
                    st.progress(
                        min(overlap_pct / 100, 1.0),
                        text=f"Turnitin overlap: {overlap_pct}%"
                    )

                    st.caption(
                        "🟢 < 15%: Aman  |  "
                        "🟡 15-30%: Perlu perhatian  |  "
                        "🔴 > 30%: Risiko tinggi"
                    )
            else:
                st.info("Humanize teks terlebih dahulu untuk melihat Turnitin Safety Score.")
        else:
            st.info("Humanize teks terlebih dahulu untuk melihat Turnitin Safety Score.")
                    
        # GPTZero Form
        st.markdown("---")
        st.markdown("#### 🕵️ GPTZero Manual Input")
        st.caption("Masukkan skor dari hasil pengecekan manual di situs GPTZero.")
        with st.form("gptzero_form"):
            st.markdown(f"**Record ID**: {st.session_state['last_record_id']}")
            col_g1, col_g2 = st.columns(2)
            with col_g1:
                gptzero_before = st.slider("Skor GPTZero (% AI) Draf Asli", 0, 100, 50, key="gptzero_before")
            with col_g2:
                gptzero_after = st.slider("Skor GPTZero (% AI) Hasil Humanize", 0, 100, 0, key="gptzero_after")
            
            submitted = st.form_submit_button("Simpan Skor Anti-Deteksi")
            if submitted:
                try:
                    resp = requests.patch(
                        f"{BACKEND_URL}/api/v1/evaluate/{st.session_state['last_record_id']}/gptzero",
                        json={
                            "gptzero_before": gptzero_before,
                            "gptzero_after": gptzero_after
                        }
                    )
                    resp.raise_for_status()
                    st.toast("✅ Skor GPTZero berhasil disimpan di riwayat!")
                    st.success("Skor GPTZero berhasil diperbarui!")
                except Exception as e:
                    st.error(f"Gagal menyimpan skor GPTZero: {e}")
    else:
        st.info("Anda belum melakukan proses Humanize. Silakan gunakan tab 'Humanize' terlebih dahulu.")


# ── Tab 3: History ────────────────────────────────────────────────────────────
with tab3:
    st.markdown("### 📜 Riwayat Evaluasi")
    st.caption("Daftar lengkap hasil proses dan evaluasi yang pernah dijalankan.")
    
    if st.button("🔄 Segarkan Data Riwayat"):
        try:
            resp = requests.get(f"{BACKEND_URL}/api/v1/evaluate/history")
            if resp.status_code == 200:
                history_data = resp.json()
                if history_data:
                    # Clean up data for dataframe
                    df = pd.DataFrame(history_data)
                    # Convert to simpler columns
                    cols_to_show = ["id", "timestamp", "style_mode", "burstiness", "ai_word_reduction", "judge_score", "gptzero_before", "gptzero_after"]
                    existing_cols = [c for c in cols_to_show if c in df.columns]
                    st.dataframe(df[existing_cols].sort_values(by="id", ascending=False), use_container_width=True)
                else:
                    st.info("Belum ada data evaluasi di database.")
        except Exception as e:
            st.error(f"Gagal mengambil riwayat: {e}")
    else:
        st.info("Klik tombol di atas untuk memuat riwayat terbaru dari database.")
